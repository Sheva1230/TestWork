#! / usr / bin / python
# -*- coding: utf-8 -*-
import collections
import sqlite3
import pandas as pd
import openpyxl
import numpy as np
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT



def wrtie_xls(array):#функция на запись в Excel(создание xlsx с помощью пандас)
    writer = pd.ExcelWriter('test.xlsx', engine='xlsxwriter')
    array.to_excel(writer, sheet_name='Factor_Table')
    writer.save()
    writer.close()

def transpon(array):#функция преобразования из словаря в табличку
        df = pd.DataFrame.from_dict(array)#делаем из словаря таблицу
        index = pd.MultiIndex.from_frame(df)#делаем мультииндекс для
        new_df = pd.DataFrame(df.unstack()).transpose()#создаем таблицу с мультииндексами и транспонируем
        return new_df

connect = sqlite3.connect('test.db')#подключение к БД
names = ['Factor','Year','World Value']#массив наименования столбцов

#Запрос с ORM(не знаю как сделать чистый запрос без ORM/курсора)

cursor = connect.execute("SELECT DISTINCT factor,year, sum(res) FROM testidprod WHERE partner IS NULL AND state IS NULL AND bs = 0 AND factor BETWEEN 1 and 2 GROUP BY factor,year ")

year_sql = connect.execute("SELECT DISTINCT year FROM testidprod WHERE partner IS NULL AND state IS NULL AND bs = 0 AND factor BETWEEN 1 and 2 GROUP BY year ")# описание запроса
itog = collections.defaultdict(dict)

array_factor = collections.defaultdict(dict)# массив для factor и year в дальнейшем будет использоваться для записи и получения суммы стобца res
'''
Array_factor - массив созданный с помощью моуля collections, defaultdict - метод который позволяет создавать словарь, который позволяет в реальном времени создать ключ словарю.  
Т.е при запросе по ключе, он будет его создовать
'''
years = []# массив годов

for row in year_sql:
    years.append(row[0])#добавления в массив

for row in cursor:
    if str(row[0]) != '' and str(row[1]) != '':
        array_factor[row[0]][row[1]] = row[2]#обавления в словарь

lenth_array = len(array_factor)# длина словаря

for key in range(lenth_array-1,lenth_array):
    for value in array_factor[lenth_array-1]:
         factor_6=array_factor[lenth_array][value]/array_factor[lenth_array-1][value]#получение значейний для 6 фактора
         array_factor[6][value] = factor_6

new_array = transpon(array_factor)
wrtie_xls(new_array)


'''В дальнейшем с помощью openpyxl, присутствует возможность 
    поправить проблему pandas  пустой строкой'''
xlsfile = 'report.xlsx'
wb = openpyxl.load_workbook(filename = 'test.xlsx')
ws = wb['Factor_Table']

row_to_delete = new_array.columns.nlevels + 1
ws.delete_rows(row_to_delete)

ws['A'+str(row_to_delete)]=new_array.index.name
for v in range(1,len(names)+1):# запись в первый столбец наименований дальнейших данных
    ws['A'+str(v)] = names[v-1]

wb.save(filename = xlsfile)
wb.close()
doc = collections.defaultdict(dict)

for value in array_factor[6]:
    doc[6][value] = array_factor[6][value]

saves_d = docx.Document()# создаем объект
saves_d.save('./report.docx')# создаем файл


docum = docx.Document('./report.docx')# открываем файл
t = docum.add_table(14, 3)#создаем таблицу в документе

counter=0

for j in range(len(names)):#заполняем заголовок таблицы
    t.cell(0,counter).text = names[j]
    counter+=1


lenth_column = len(t.columns[1].cells)
first_cell = t.cell(1,0)
last_cell = t.cell(lenth_column-1,0)
merge_Cell = first_cell.merge(last_cell)


for key in doc:
    t.cell(1,0).text=str(key)
    for i in range(0,len(years)):
        t.cell(i+1,1).text= str(years[i])
        t.cell(i+1,2).text =str(doc[key][years[i]])

t.cell(1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
t.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

docum.add_paragraph("")
# save the doc
docum.save('./test.docx')

